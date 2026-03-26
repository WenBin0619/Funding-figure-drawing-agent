from docx import Document


def read_word_document(file_path: str) -> str:
    document = Document(file_path)
    
    full_text = []
    
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            full_text.append(text)
    
    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n\n".join(full_text)


def read_word_document_with_structure(file_path: str) -> dict:
    document = Document(file_path)
    
    result = {
        "paragraphs": [],
        "tables": [],
        "full_text": ""
    }
    
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            result["paragraphs"].append({
                "text": text,
                "style": paragraph.style.name if paragraph.style else None
            })
    
    for table_idx, table in enumerate(document.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append({
            "index": table_idx,
            "data": table_data
        })
    
    result["full_text"] = read_word_document(file_path)
    
    return result


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python word_reader.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    content = read_word_document(file_path)
    print(content)
